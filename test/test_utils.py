from src.utils import extract_text_from_file, generate_word_list, convert_word_list_to_csv_with_translations, check_for_new_words, get_user_language, convert_word_list_to_csv, extract_file_list, extract_text_from_url, extract_text_from_mkv 
import pytest
import csv
import docx
from reportlab.pdfgen.canvas import Canvas
from unittest.mock import patch, Mock
import requests
from ebooklib import epub
from pathlib import Path


@pytest.fixture
def example_srt(tmp_path):
    """Creates a temporary .srt file for testing."""
    example_srt = tmp_path / 'example.srt'
    return example_srt

@pytest.fixture
def example_csv(tmp_path):
    """Creates a temporary .csv file for testing."""
    example_csv = tmp_path / 'example.csv'
    return example_csv

@pytest.fixture
def example_text():
    """Returns an example text block for test cases."""
    example_text = '''Lorem ipsum dolor sit amet, consectetur adipiscing elit. Vivamus non eros id mi volutpat faucibus. Donec id nulla non nunc iaculis egestas id vestibulum ligula. Duis auctor massa cursus volutpat venenatis. In hac habitasse platea dictumst. Aliquam sit amet laoreet dui, lobortis fermentum neque. Ut quis semper massa, a blandit quam. Duis in laoreet quam, vel facilisis orci. Aliquam at lorem non ligula gravida rhoncus. In eu suscipit lectus. Suspendisse potenti. Duis porta libero orci, id lacinia libero finibus sit amet. Donec nec magna ut ex hendrerit suscipit. In enim quam, aliquam a orci quis, volutpat vestibulum urna. Integer euismod nec diam ac congue. In sit amet sem tortor.'''
    return example_text

@pytest.fixture
def file_extensions():
    """Creates a list of valid file extensions."""
    return ['.srt', '.txt', '.md']

class TestExtractTextFromFile:
    """Tests for the extract_text_from_file() function in utils.py"""

    def test_removes_timestamps(self, example_srt):
        """Ensures that timestamps are correctly removed from SRT files."""
        example_srt.write_text("""1
                               00:00:35,077 --> 00:00:36,203
                               Hello!\n\n2
                               00:00:41,291 --> 00:00:42,751
                               Hello!""")
        expected = '''Hello!\n\nHello!'''
        output = extract_text_from_file(example_srt)
        assert output == expected
    
    def test_removes_italic_markers(self, example_srt):
        """Ensures that italic markers are correctly removed from SRT files."""
        example_srt.write_text('''5
                               00:04:30,604 --> 00:04:32,231
                               <i>Good morning, Refiners.</i>''')
        expected = '''Good morning, Refiners.'''
        output = extract_text_from_file(example_srt)
        assert output == expected
    
    def test_handles_multiple_lines_of_text(self, example_srt):
        """Checks that multiline text remains properly formatted after processing."""
        example_srt.write_text('''6
                                00:04:33,232 --> 00:04:36,359
                                <i>This is Mr. Milchick from work,\nand I'm thrilled to welcome you</i>''')
        expected = '''This is Mr. Milchick from work,\nand I'm thrilled to welcome you'''
        output = extract_text_from_file(example_srt)
        assert output == expected
    
    def test_removes_any_other_html_tags(self, example_srt):
        """Ensures that all HTML tags are correctly removed from SRT files."""
        example_srt.write_text('''5
                               00:04:30,604 --> 00:04:32,231
                               <b>Good morning, Refiners.</b>''')
        expected = '''Good morning, Refiners.'''
        output = extract_text_from_file(example_srt)
        assert output == expected
    
    def test_text_unchanged_if_no_timestamps_or_html_tags(self, example_srt, example_text):
        """Checks that text is returned unchanged if there are no timestamps or HTML tags."""
        example_srt.write_text(example_text)
        output = extract_text_from_file(example_srt)
        assert example_srt.read_text() == example_text
        assert output == example_text
    
    def test_error_message_returned_if_filepath_invalid(self):
        """Ensures that a FileNotFoundError is raised for nonexistent files."""
        with pytest.raises(FileNotFoundError) as err:
            extract_text_from_file('example1.srt')
        assert str(err.value) == "Error: The file 'example1.srt' was not found."
    
    def test_IO_error_occurs_if_file_format_is_invalid(self, tmp_path):
        """Tests that an IOError is raised when an unsupported file format is used."""
        example_mkv = tmp_path / 'example.mkv'
        example_mkv.touch()
        with pytest.raises(IOError) as err:
            extract_text_from_file(example_mkv)
        assert str(err.value) == "Error: Could not read the file contents of 'example.mkv'. File format is invalid."

    def test_blank_document_remains_unchanged(self, example_srt):
        """Checks that a file with no content remains unchanged."""
        example_srt.touch()
        output = extract_text_from_file(example_srt)
        assert not example_srt.read_text()
        assert not output
    
    def test_handles_md_files(self, tmp_path):
        """Checks that md file can be successfully processed."""
        example_md = tmp_path / 'example.md'
        example_md.write_text("example text")
        assert extract_text_from_file(example_md) == "example text"
    
    def test_handles_docx_files(self, tmp_path):
        """Checks that docx files can be successfully processed."""
        example_docx = tmp_path / 'example.docx'
        doc = docx.Document()
        doc.add_paragraph("here is some text")
        doc.save(example_docx)
        assert extract_text_from_file(example_docx) == "here is some text"
    
    def test_removes_unwanted_formatting_from_docx_files(self, tmp_path):
        """Checks that unwanted formatting artifacts are removed from docx files."""
        example_docx = tmp_path / 'example.docx'
        doc = docx.Document()
        doc.add_paragraph(r"here is \an8}some text")
        doc.save(example_docx)
        assert extract_text_from_file(example_docx) == "here is some text"
    
    def test_handles_pdf_files(self, tmp_path):
        """Checks that pdf files can be successfully processed."""

        def create_pdf(pdf_path, text):
            file = Canvas(str(pdf_path))
            file.drawString(72, 72, text)
            file.save()

        example_text = "here is some text"
        test_pdf = tmp_path / 'example.pdf'
        create_pdf(test_pdf, example_text)

        assert example_text in extract_text_from_file(test_pdf)
    
    def test_handles_epub_files(self, tmp_path):
        """Checks that epub files can be successfully processed."""

        epub_path = tmp_path / "example.epub"

        test_epub = epub.EpubBook()
        chapter1 = epub.EpubHtml(title='Chapter 1', file_name="chap1.xhtml")
        chapter1.content = "<h1>Chapter 1</h1><p>here is some text</p>"

        chapter2 = epub.EpubHtml(title='Chapter 2', file_name="chap2.xhtml")
        chapter2.content = "<h1>Chapter 2</h1><p>and now some more text</p>"

        test_epub.add_item(chapter1)
        test_epub.add_item(chapter2)

        test_epub.spine = ["nav", chapter1, chapter2]
        test_epub.add_item(epub.EpubNcx())
        test_epub.add_item(epub.EpubNav())

        epub.write_epub(str(epub_path), test_epub)

        assert "here is some text" in extract_text_from_file(epub_path)       
        assert "and now some more text" in extract_text_from_file(epub_path)


class TestGenerateWordList:
    """Tests for the generate_word_list() function in utils.py."""

    def test_empty_dict_returned_if_no_input(self):
        """Should return an empty dictionary if passed an empty string."""
        assert generate_word_list('') == {}

    def test_returns_dictionary(self):
        """Should return a dictionary when given valid input."""
        output = generate_word_list('hello')
        assert isinstance(output, dict)

    def test_counts_single_word(self):
        """Should count occurrences of a single word correctly."""
        output = generate_word_list('hello')
        assert output == {'hello': 1}
    
    def test_counts_multiple_words(self):
        """Should count occurrences of multiple words correctly."""
        assert generate_word_list('hello world') == {'hello': 1, 'world': 1}
    
    def test_counts_multiple_instances_of_same_word(self):
        """Should correctly count multiple occurrences of the same word."""
        assert generate_word_list('hello world hello') == {'hello': 2, 'world': 1}

    def test_ignores_capital_letters(self):
        """Should treat uppercase and lowercase words the same."""
        assert generate_word_list('Hello world hello') == {'hello': 2, 'world': 1}

    def test_filters_out_punctuation(self):
        """Should ignore punctuation characters."""
        text = "The quick brown fox jumps over the lazy dog. The dog was not amused?"
        assert generate_word_list(text) == {'amused': 1, 'brown': 1, 'dog': 2, 'fox': 1, 'jumps': 1, 'lazy': 1, 'not': 1, 'over': 1, 'quick': 1, 'the': 3, 'was': 1}
    
    def test_filters_out_additional_punctuation_characters(self):
        """Should ignore rarer punctuation characters such as tags or Spanish-specific question marks."""
        text = '''¿Sueles leer «antes» de dormir? Al principio: <Si trabajas duro, conseguirás lo que quieres.>>'''
        assert generate_word_list(text) == {'sueles': 1, 'leer': 1, 'antes': 1, 'de': 1, 'dormir': 1, 'al': 1, 'principio': 1, 'si': 1, 'trabajas': 1, 'duro': 1, 'conseguirás': 1, 'lo': 1, 'que': 1, 'quieres': 1}
    
    def test_handles_text_in_cyrillic_script(self):
        """Should correctly count words in Cyrillic script."""
        text = 'Старик был сварливым'
        assert generate_word_list(text) == {'старик': 1, 'был': 1, 'сварливым': 1}
    
    def test_ignores_whitespace_characters(self):
        """Should ignore newline or tab characters."""
        text = '''Hello\nworld\teverything\nis\tfine'''
        assert generate_word_list(text) == {'hello': 1, 'world': 1, 'everything': 1, 'is': 1, 'fine': 1}
    
    def test_hyphens_ignored_in_middle_of_words(self):
        """Should retain hyphens in compound words."""
        text = '''hello, first, second, físico-químico.'''
        assert generate_word_list(text) == {'hello': 1, 'first': 1, 'second': 1, 'físico-químico': 1}
    
    def test_removes_numbers_from_word_list(self):
        """Should remove numbers from the text."""
        text = "4 foxes jumped over 2 3-year-old dogs?"
        assert generate_word_list(text) == {'foxes': 1, 'jumped': 1, 'over': 1, '3-year-old': 1, 'dogs': 1}


class TestConvertToCSVWithTranslations:
    """Tests for the convert_word_list_to_csv_with_translations() function in utils.py."""

    def test_creates_csv_file(self, example_csv):
        """Should create a CSV file from the word list."""
        input = {'hello': 1}
        convert_word_list_to_csv_with_translations(input, example_csv, "en", "de")
        assert example_csv.exists()
    
    def test_converts_single_key_value_pair_to_csv(self, example_csv):
        """Should correctly write a single word-frequency pair to the CSV file."""
        input = {'hello': 1}
        convert_word_list_to_csv_with_translations(input, example_csv, "en", "de")
        with open(example_csv, newline="") as file:
            reader = csv.reader(file)
            first_row = next(reader)
            assert first_row[0] == 'hello: 1'
    
    def test_sorts_and_converts_multiple_key_value_pairs(self, example_csv):
        """Should sort words alphabetically and write them to the CSV file."""
        input = {'hello': 1, 'world': 1, 'abacus': 1}
        convert_word_list_to_csv_with_translations(input, example_csv, "en", "de")
        with open(example_csv, newline="") as file:
            reader = csv.reader(file)
            assert next(reader)[0] == 'abacus: 1'
            assert next(reader)[0] == 'hello: 1'
            assert next(reader)[0] == 'world: 1'

    def test_adds_translation_to_single_key_value_pair(self, example_csv):
        """Should correctly add translations for a single word."""
        input = {'hello': 1}
        convert_word_list_to_csv_with_translations(input, example_csv, "en", "fr")
        with open(example_csv, newline="") as file:
            reader = csv.reader(file)
            first_row = next(reader)
            assert first_row[0] == 'hello: 1'
            assert first_row[1].lower() == 'bonjour'
    
    def test_adds_translations_for_multiple_words(self, example_csv):
        """Should correctly add translations for multiple words."""
        input = {'hello': 1, 'world': 1, 'abacus': 1}
        convert_word_list_to_csv_with_translations(input, example_csv, "en", "fr")
        with open(example_csv, newline="") as file:
            reader = csv.reader(file)
            first_row = next(reader)
            assert first_row[0] == 'abacus: 1'
            assert first_row[1].lower() == 'abaque'
            second_row = next(reader)
            assert second_row[0] == 'hello: 1'
            assert second_row[1].lower() == 'bonjour'
            third_row = next(reader)
            assert third_row[0] == 'world: 1'
            assert third_row[1].lower() == 'monde'

class TestCheckForNewWords:
    """Tests for the check_for_new_words() function."""

    def test_returns_new_dict(self):
        """Should return a new dictionary object instead of modifying the original."""
        input_dict = {'hello': 1}
        input_set = {'word'}
        output = check_for_new_words(input_dict, input_set)
        assert output is not input_dict
    
    def test_dict_unchanged_if_set_contains_no_matches(self):
        """Should return the same dictionary if no words from the set match."""
        input_dict = {'hello': 1}
        input_set = {'word'}
        output = check_for_new_words(input_dict, input_set)
        assert output == {'hello': 1}
    
    def test_dict_item_removed_if_match_found(self):
        """Should remove matching words from the dictionary."""
        input_dict = {'hello': 1, 'world': 1}
        input_set = {'world'}
        assert check_for_new_words(input_dict, input_set) == {'hello': 1}

class TestGetUserLanguage:
    """Tests for the get_user_language() function."""

    def test_returns_language_code_entered_by_user(self):
        """Should return the same language code if a valid code is entered."""
        assert get_user_language(['es']) == 'es'
    
    def test_returns_language_code_if_user_enters_language_name(self):
        """Should return the same language code if a valid code is entered."""
        assert get_user_language(['spanish']) == 'es'
    
    def test_ignores_capitalisation(self):
        """Should handle capitalised or uppercase language names correctly."""
        assert get_user_language(['Spanish']) == 'es'
        assert get_user_language(['ES']) == 'es'
        assert get_user_language(['SPANISH']) == 'es'
    
    def test_handles_invalid_input_followed_by_valid_input(self, capsys):
        """Should prompt again if invalid input is given, then accept valid input."""
        assert get_user_language(['invalid', 'French']) == 'fr'
        captured = capsys.readouterr()
        assert "Invalid input." in captured.out
    
    def test_available_languages_printed_upon_user_request(self, capsys):
        """Should print a list of available languages when requested."""
        assert get_user_language(['l', 'French']) == 'fr'
        captured = capsys.readouterr()
        assert "Available languages: " in captured.out
        assert 'french: fr' in captured.out
        assert 'spanish: es' in captured.out
    

class TestConvertToCSV:
    """Tests for the convert_word_list_to_csv() function."""
    
    def test_creates_csv_file(self, example_csv):
        """Should create a CSV file at the specified path."""
        input = {'hello': 1}
        convert_word_list_to_csv(input, example_csv)
        assert example_csv.exists()
    
    def test_converts_single_key_value_pair_to_csv(self, example_csv):
        """Should correctly write a single word-count pair to the CSV file."""
        input = {'hello': 1}
        convert_word_list_to_csv(input, example_csv)
        with open(example_csv, newline="") as file:
            reader = csv.reader(file)
            first_row = next(reader)
            assert first_row[0].lstrip('\ufeff') == 'hello'
            assert int(first_row[1]) == 1
    
    def test_sorts_and_converts_multiple_key_value_pairs(self, example_csv):
        """Should correctly write and sort multiple word-count pairs in the CSV file."""
        input = {'hello': 1, 'world': 1, 'abacus': 1}
        convert_word_list_to_csv(input, example_csv)
        with open(example_csv, newline="") as file:
            reader = csv.reader(file)
            first_row = next(reader)
            assert first_row[0].lstrip('\ufeff') == 'abacus'
            assert int(first_row[1]) == 1
            second_row = next(reader)
            assert second_row[0] == 'hello'
            assert int(second_row[1]) == 1
            third_row = next(reader)
            assert third_row[0] == 'world'
            assert int(third_row[1]) == 1
    
class TestExtractFileList:

    def test_extracts_single_file_from_directory(self, file_extensions, tmp_path):
        """Should correctly extract a valid file from a directory."""
        file = tmp_path / "file.txt"
        file.write_text("test")
        assert extract_file_list(tmp_path, file_extensions) == [file]

    def test_extracts_multiple_files_from_directory(self, file_extensions, tmp_path):
        """Should correctly extract valid files from a directory."""
        file1 = tmp_path / "file1.txt"
        file1.write_text("test1")
        file2 = tmp_path / "file2.srt"
        file2.write_text("test2")
        assert extract_file_list(tmp_path, file_extensions) == [file1, file2]
    
    def test_ignores_files_with_invalid_formats(self, file_extensions, tmp_path):
        """Should return an empty list if file formats are invalid."""
        file1 = tmp_path / "file1.pdf"
        file1.write_text("test")
        assert extract_file_list(tmp_path, file_extensions) == []
    
    def test_handles_mixture_of_valid_and_invalid_files(self, file_extensions, tmp_path):
        """Should save valid files to the list while ignoring invalid files."""
        file1 = tmp_path / "file1.pdf"
        file1.write_text("test1")
        file2 = tmp_path / "file2.srt"
        file2.write_text("test2")
        assert extract_file_list(tmp_path, file_extensions) == [file2]

    def test_raises_error_if_invalid_directory(self, file_extensions):
        """Should raise error if an invalid directory is given."""
        with pytest.raises(ValueError) as err:
            extract_file_list("hello", file_extensions)
        assert str(err.value) == f"Invalid directory: hello"
    
    def test_locates_file_in_subfolder(self, file_extensions, tmp_path):
        """Should correctly identify file located in a subfolder of the given directory."""
        subfolder = tmp_path / "files"
        subfolder.mkdir()
        file1 = subfolder / "test_file.txt"
        file1.write_text("test")
        assert extract_file_list(tmp_path, file_extensions)

@pytest.fixture
def mock_get_request():
    """Creates a test response body."""
    with patch("requests.get") as mock_get:
        mock_response = Mock()
        mock_response.content = b"""
        <html>
            <body>
                <article class="ssrcss-z9afcx-ArticleWrapper e1nh2i2l3">
                    <header data-component="headline-block" class="ssrcss-bwbna7-ComponentWrapper-HeadlineComponentWrapper egtrm1f0">
                        <h1 id="main-heading" type="headline" tabindex="-1" class="ssrcss-1s9pby4-Heading e10rt3ze0">
                            <span role="text">
                                Reeves says UK beginning to turn corner as growth beats forecasts
                            </span>
                        </h1>
                    </header>
                    <div class="ssrcss-1w03aro-RichTextComponentWrapper ep2nwvo0">
                        <p class="ssrcss-1q0x1qg-Paragraph ejhz7w10">
                            <b>The growth figure was stronger than</b>
                        </p>
                        <p class="ssrcss-1q0x1qg-Paragraph ejhz7w10">
                            Liberal Democrat Treasury spokesperson
                        </p>
                        <p class="ssrcss-1q0x1qg-Paragraph ejhz7w10">
                            predictions are highly volatile
                        </p>
                    </div>
                </article>
            </body>
        </html>
        """
        mock_get.return_value = mock_response
        yield mock_get

@pytest.fixture
def mock_error():
    """Mocks requests.get to raise an error."""
    with patch("requests.get") as mock_error:
        mock_error.side_effect = requests.exceptions.RequestException
        yield mock_error


class TestExtractTextFromUrl:

    """Tests for the extract_text_from_url function."""

    def test_extracts_text_from_url(self, mock_get_request):
        """Checks body and header are successfully extracted from URL."""
        output = extract_text_from_url("test")
        assert "The growth figure was stronger than" in output
        assert "Liberal Democrat Treasury spokesperson" in output
        assert "predictions are highly volatile" in output
        assert "Reeves says UK beginning to turn corner as growth beats forecasts" in output

    def test_raises_value_error_if_request_fails(self, mock_error):
        """Checks whether error is raised if URL is invalid."""

        with pytest.raises(ValueError) as err:
            extract_text_from_url("www.invalid-url.com")
        assert str(err.value) == "Text extraction failed. URL may be invalid."        


@pytest.fixture
def mock_mkv_subs(tmp_path):
    """Creates a mock subprocess on an .srt file."""
    example_srt = tmp_path / "example.srt"
    example_srt.write_text('''6
                                00:04:33,232 --> 00:04:36,359
                                <i>This is Mr. Milchick from work,\nand I'm thrilled to welcome you</i>''')

    with patch("tempfile.NamedTemporaryFile") as mock_temp, \
        patch("subprocess.run") as mock_subp:
        mock_temp.return_value.__enter__.return_value.name = str(example_srt)
        yield {
            "mock_temp": mock_temp,
            "mock_subp": mock_subp,
            "temp_path": str(example_srt)
            }

class TestExtractTextFromMkv:
    """Tests for the extract_text_from_mkv function."""

    def test_extracts_text_correctly(self, mock_mkv_subs):
        """Checks that text is successfully extracted by the function."""
        expected = '''This is Mr. Milchick from work,\nand I'm thrilled to welcome you'''

        assert extract_text_from_mkv("test.mkv", 2) == expected

    def test_calls_correct_subprocess_arguments(self, mock_mkv_subs):
        """Checks that the correct subprocess arguments are used."""
        extract_text_from_mkv("test.mkv", 2)

        mock_mkv_subs["mock_subp"].assert_called_once_with([
            "mkvextract",
            "test.mkv",
            "tracks",
            f"2:{mock_mkv_subs['temp_path']}"
        ], check=True)

    def test_temp_file_removed_after_function_complete(self, mock_mkv_subs):
        """Checks that the temporary .txt file is deleted."""
        extract_text_from_mkv("test.mkv", 2)

        assert not Path(mock_mkv_subs["temp_path"]).exists()